"""
Renders an ApprovalNote to a .docx file that replicates the layout of the
reference sample document:
  - Font: Book Antiqua, 12pt (w:sz 24) throughout
  - Header line: "CDRSL/BOND/.../.../.." left, date right (tab-aligned)
  - Bold underlined Subject line
  - "Para 1" / "Para 2" / "Para 4" bold labels followed by normal text
  - A 6-column table (Sl No | Consignment Charges | Invoice/Ref | Amount | CHA Ref | SYS Ref)
    with a bold header row and a bold Total row, using the exact column widths
    (in twips) measured from the sample: 554, 2569, 2482, 1583, 1400, 1352
  - Signature block with 5 fixed signatory lines
"""
import os
from docx import Document
from docx.shared import Pt, Twips, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "Book Antiqua"
FONT_SIZE = Pt(12)
COL_WIDTHS_TWIPS = [554, 2569, 2482, 1583, 1400, 1352]  # measured from sample document.xml

SIGNATORIES = [
    "Sr. Superintendent Duty Free",
    "Sr. Superintendent SG Duty Free",
    "DGM (CDRSL)                                               AGM (FIN)",
    "CFO CDRSL",
    "MD, CDRSL",
]


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            el = OxmlElement(tag)
            for key, val in edge_data.items():
                el.set(qn(f'w:{key}'), str(val))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _style_run(run, bold=False, underline=False, size=None):
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE
    run.bold = bold
    run.underline = underline
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)


def _add_paragraph(doc, spacing_line=240, jc=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = spacing_line / 240.0
    if jc == "both":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def render_approval_note_docx(note, shipment, invoices, output_path):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Twips(11906)   # A4
    section.page_height = Twips(16838)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # ---- Header line: document number (left) ... date (right) ----
    header_p = doc.add_paragraph()
    header_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT)
    r1 = header_p.add_run(note.document_number)
    _style_run(r1, bold=True)
    r2 = header_p.add_run("\t" + note.document_date.strftime("%d/%m/%Y"))
    _style_run(r2, bold=True)

    doc.add_paragraph()  # spacer

    # ---- Subject line (bold, underlined) ----
    subj_p = _add_paragraph(doc, spacing_line=360)
    r = subj_p.add_run("Sub: Payment of Clearance Charges to ")
    _style_run(r, bold=True, underline=True)
    r = subj_p.add_run(f"M/s {shipment.cha_name} ")
    _style_run(r, bold=True, underline=True)
    r = subj_p.add_run("\u2013 our CHA. [")
    _style_run(r, bold=True, underline=True)
    system_bits = []
    if shipment.grn_number:
        system_bits.append(f"GRN-{shipment.grn_number}")
    if shipment.boe_number:
        system_bits.append(f"BOE-{shipment.boe_number}")
    if shipment.supplier_invoice_ref:
        system_bits.append(f"Supplier Invoice-{shipment.supplier_invoice_ref}")
    r = subj_p.add_run("SYSTEM " + " ".join(system_bits))
    _style_run(r, bold=True, underline=True)

    doc.add_paragraph()
    doc.add_paragraph()

    # ---- Para 1 ----
    p1 = _add_paragraph(doc, spacing_line=240, jc="both")
    r = p1.add_run("Para 1")
    _style_run(r, bold=True)
    r = p1.add_run(": ")
    _style_run(r)
    r = p1.add_run(note.paragraph_1 or "")
    _style_run(r)

    # ---- Table ----
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["Sl No", "Consignment Charges", "Invoice /Ref:", "Amount", "CHA Ref", "SYS Ref"]
    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr_cells[i].width = Twips(COL_WIDTHS_TWIPS[i])
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(htext)
        _style_run(r, bold=True)

    total_amount = 0
    for idx, inv in enumerate(invoices, start=1):
        row = table.add_row().cells
        row[0].width = Twips(COL_WIDTHS_TWIPS[0])
        row[1].width = Twips(COL_WIDTHS_TWIPS[1])
        row[2].width = Twips(COL_WIDTHS_TWIPS[2])
        row[3].width = Twips(COL_WIDTHS_TWIPS[3])
        row[4].width = Twips(COL_WIDTHS_TWIPS[4])
        row[5].width = Twips(COL_WIDTHS_TWIPS[5])

        r = row[0].paragraphs[0].add_run(f"{idx}.")
        _style_run(r)

        vendor = f" ({inv.vendor_name})" if inv.vendor_name else ""
        r = row[1].paragraphs[0].add_run(f"{inv.short_label()}{vendor}")
        _style_run(r)

        r = row[2].paragraphs[0].add_run(inv.invoice_number or "")
        _style_run(r)

        amt = float(inv.invoice_amount or 0)
        total_amount += amt
        r = row[3].paragraphs[0].add_run(f"\u20b9{amt:,.2f}")
        _style_run(r)

        r = row[4].paragraphs[0].add_run(inv.cha_ref or "")
        _style_run(r)

        r = row[5].paragraphs[0].add_run(inv.srn or "")
        _style_run(r)

    # Total row
    total_row = table.add_row().cells
    for i in range(4):
        total_row[i].width = Twips(COL_WIDTHS_TWIPS[i])
    total_row[4].width = Twips(COL_WIDTHS_TWIPS[4])
    total_row[5].width = Twips(COL_WIDTHS_TWIPS[5])
    r = total_row[2].paragraphs[0].add_run("Total")
    _style_run(r, bold=True)
    r = total_row[3].paragraphs[0].add_run(f"\u20b9{total_amount:,.2f}")
    _style_run(r, bold=True)

    doc.add_paragraph()

    # ---- Para 2 (recoverable explanation) ----
    if note.paragraph_2:
        p2 = _add_paragraph(doc, spacing_line=240, jc="both")
        r = p2.add_run("Para 2: ")
        _style_run(r, bold=True)
        r = p2.add_run(note.paragraph_2)
        _style_run(r)

    # ---- Para 3 (optional remarks) ----
    if note.paragraph_3:
        p3 = _add_paragraph(doc, spacing_line=240, jc="both")
        r = p3.add_run("Para 3: ")
        _style_run(r, bold=True)
        r = p3.add_run(note.paragraph_3)
        _style_run(r)

    # ---- Para 4 ----
    p4 = _add_paragraph(doc, spacing_line=240, jc="both")
    r = p4.add_run("Para 4: ")
    _style_run(r, bold=True)
    r = p4.add_run(note.paragraph_4 or "")
    _style_run(r)

    # ---- Signature block ----
    doc.add_paragraph()
    doc.add_paragraph()
    for i, sig in enumerate(SIGNATORIES):
        sp = _add_paragraph(doc, spacing_line=240)
        r = sp.add_run(sig)
        _style_run(r)
        if i < len(SIGNATORIES) - 1:
            doc.add_paragraph()
            doc.add_paragraph()

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
